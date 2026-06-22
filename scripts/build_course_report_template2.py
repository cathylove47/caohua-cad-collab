from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = Path("/Users/cathy/Downloads/课程设计报告模板2.docx")
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "课程设计报告-模板2完成版.docx"

FIGURES = ROOT / "report-latex" / "figures"
NUMBERING_IDS = list(range(1, 39))


def set_run_font(run, size: float = 12, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def format_body_paragraph(paragraph: Paragraph, first_line_indent: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    if first_line_indent:
        fmt.first_line_indent = Cm(0.74)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def format_list_paragraph(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def clear_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_table(table: Table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    size: float = 12,
    bold: bool = False,
    first_line_indent: bool = True,
    center: bool = False,
    color: RGBColor | None = None,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        format_body_paragraph(paragraph, first_line_indent=first_line_indent)


def add_body(doc: DocumentObject, text: str) -> Paragraph:
    paragraph = doc.add_paragraph()
    format_body_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def next_numbering_id() -> int:
    if not NUMBERING_IDS:
        raise RuntimeError("No numbering definitions left in template.")
    return NUMBERING_IDS.pop(0)


def apply_numbering(paragraph: Paragraph, num_id: int) -> None:
    paragraph.style = "List Paragraph"
    p_pr = paragraph._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    p_pr.append(num_pr)


def add_bullets(doc: DocumentObject, items: list[str]) -> None:
    num_id = next_numbering_id()
    for item in items:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, num_id)
        format_list_paragraph(paragraph)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: DocumentObject, items: list[str]) -> None:
    num_id = next_numbering_id()
    for item in items:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, num_id)
        format_list_paragraph(paragraph)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_heading_1(doc: DocumentObject, text: str, page_break: bool = True) -> None:
    if page_break:
        doc.add_page_break()
    paragraph = doc.add_heading(text, level=1)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    for run in paragraph.runs:
        set_run_font(run, size=16, bold=True)


def add_heading_2(doc: DocumentObject, text: str) -> None:
    paragraph = doc.add_heading(text, level=2)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=14, bold=True)


def add_caption(doc: DocumentObject, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(9)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)


def add_reference_list(doc: DocumentObject, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.first_line_indent = Cm(-0.5)
        run = paragraph.add_run(f"{idx}. {item}")
        set_run_font(run)


def add_figure(doc: DocumentObject, image_path: Path, caption: str, width_inches: float = 5.9) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    add_caption(doc, caption)


def set_cell_text(cell: _Cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=bold)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: DocumentObject, rows: list[list[str]], widths_cm: list[float]) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), value, bold=row_idx == 0, center=col_idx != 2)
    header_row = table.rows[0]
    tr_pr = header_row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    doc.add_paragraph()
    return table


def add_footer_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph(paragraph)

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=10.5)


def strip_template_body(doc: DocumentObject) -> None:
    for table in list(doc.tables):
        delete_table(table)
    break_index = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if any(child.tag == qn("w:r") and child.find(qn("w:br")) is not None for child in paragraph._element):
            break_index = idx
            break
    for paragraph in list(doc.paragraphs)[break_index + 1 :][::-1]:
        delete_paragraph(paragraph)


def configure_cover(doc: DocumentObject) -> None:
    paragraphs = doc.paragraphs
    page_break_paragraph = paragraphs[14]
    date_paragraph = paragraphs[9]
    set_paragraph_text(
        paragraphs[1],
        "设计题目：基于云的协同机械 CAD 系统设计与实现",
        size=16,
        bold=False,
        first_line_indent=False,
    )
    set_paragraph_text(paragraphs[3], "姓　　名：待填写", size=14, bold=False, first_line_indent=False)
    set_paragraph_text(paragraphs[4], "学　　号：待填写", size=14, bold=False, first_line_indent=False)
    set_paragraph_text(paragraphs[5], "专　　业：待填写", size=14, bold=False, first_line_indent=False)
    set_paragraph_text(paragraphs[6], "班　　级：待填写", size=14, bold=False, first_line_indent=False)
    set_paragraph_text(paragraphs[7], "指导教师：待填写", size=14, bold=False, first_line_indent=False)
    for paragraph in [paragraphs[13], paragraphs[12], paragraphs[11], paragraphs[10]]:
        delete_paragraph(paragraph)
    today = date.today()
    set_paragraph_text(
        date_paragraph,
        f"完成日期：{today.year}年{today.month:02d}月{today.day:02d}日",
        size=14,
        bold=False,
        first_line_indent=False,
    )
    clear_paragraph(page_break_paragraph)
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def build_report() -> Path:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE_PATH))
    doc.core_properties.title = "工业软件开发技术课程设计报告"
    doc.core_properties.subject = "基于云的协同机械 CAD 系统设计与实现"
    doc.core_properties.author = "待填写"

    configure_cover(doc)
    strip_template_body(doc)

    add_heading_1(doc, "1 总体架构设计", page_break=False)
    add_heading_2(doc, "1.1 课程设计目标与实现范围")
    add_body(
        doc,
        "本课程设计的目标不是构建完整工业级 CAD 平台，而是在本地环境中实现一个稳定可演示、功能闭环清晰、便于课程答辩说明的协同机械 CAD MVP。系统以“一个浏览器端 Project + 一个协同后端 + 一套本地数据文件”的思路展开，重点完成参数化建模、浏览器渲染显示、项目保存和多客户端协同设计。",
    )
    add_body(
        doc,
        "在实现范围上，前端负责新建 Project、创建基础参数化实体、查看与编辑对象参数、进行视图旋转、拖动、缩放与拉伸操作以及参与协同编辑；后端负责保存项目和实体数据文件，并对同一项目下多个客户端的操作进行同步广播。老师模板中关于微服务、复杂格式转换、装配、FEA 和工业级几何内核的内容，均不作为本课程设计的必做实现。",
    )
    add_heading_2(doc, "1.2 系统总体架构")
    add_body(
        doc,
        "系统总体采用“前端浏览器建模与渲染 + 后端房间协同与持久化”的二层主架构。前端通过 React、Three.js 和 Zustand 负责界面与交互逻辑，后端通过 Express、WebSocket 和本地 JSON 文件负责项目状态维护、项目保存和多人协同。这样的设计既符合课程设计对演示稳定性的要求，也更容易解释系统职责分工。",
    )
    add_bullets(
        doc,
        [
            "前端浏览器端：负责 Project 创建、参数化实体建模、三维渲染与对象编辑。",
            "后端协同服务：负责项目状态维护、在线用户管理、操作广播与版本快照。",
            "本地数据层：负责保存项目文件和实体数据文件，为 Load 和 Restore 提供依据。",
        ],
    )
    add_heading_2(doc, "1.3 技术栈选型")
    add_table(
        doc,
        [
            ["层次", "技术选型", "技术选型作用"],
            ["前端 UI 框架", "React + TypeScript", "构建组件化 CAD 交互界面"],
            ["前端构建工具", "Vite", "提供快速开发与打包构建能力"],
            ["3D 渲染引擎", "Three.js", "完成几何体渲染、相机与场景管理"],
            ["前端状态管理", "Zustand", "管理对象状态、历史栈与协同状态"],
            ["后端服务框架", "Node.js + Express", "提供 HTTP API 与服务入口"],
            ["实时协同通信", "WebSocket + ws", "完成房间广播、在线用户与远端光标同步"],
            ["数据持久化", "本地 JSON 文件", "保存房间模型状态与版本快照"],
            ["几何内核扩展", "OpenCASCADE/Wasm（预留）", "作为后续工业级几何计算扩展方向"],
        ],
        [2.8, 4.6, 9.3],
    )
    add_body(
        doc,
        "上述技术栈优先满足本地可运行、依赖少、调试方便和课程答辩易讲解四个要求。其中后端采用单体服务而不是微服务，是有意的工程取舍：当前规模下单服务结构更直观、更稳定，也更贴合课程版 MVP 的目标。",
    )
    add_heading_2(doc, "1.4 协同与数据流")
    add_body(
        doc,
        "用户在登录页输入用户名和 roomId 后，即可新建或进入一个 Project。前端先加载该 Project 最近一次保存状态，再通过 WebSocket 加入同名协同房间。之后用户每次新增实体、修改参数、保存项目或恢复版本，都会在前端本地更新，并由后端广播给其他客户端。",
    )
    add_figure(doc, FIGURES / "collab-dataflow.png", "图 1-2 项目协同与持久化数据流图", width_inches=4.9)

    add_heading_1(doc, "2 前端系统设计")
    add_heading_2(doc, "2.1 前端总体结构")
    add_body(
        doc,
        "前端主要由登录页、CAD 工作区、Three.js 视图层和 Zustand 状态层组成。登录页负责输入用户名和 roomId；工作区负责布局工具栏、特征树、属性面板和状态区；Three.js 视图层负责三维几何体渲染；状态层负责对象数据、Project 状态、历史记录和协同连接状态。",
    )
    add_figure(doc, FIGURES / "module-uml.png", "图 2-1 前后端核心模块关系 UML 图", width_inches=5.4)
    add_heading_2(doc, "2.2 Project 创建与 UI 布局")
    add_body(
        doc,
        "在前端 UI 中，用户通过输入用户名和 roomId 来新建或进入一个 Project。系统把 roomId 作为课程版 Project 的唯一标识，既简化了登录流程，也方便演示协同。进入工作区后，页面采用典型 CAD 布局：顶部是 Ribbon 风格工具栏，左侧是 Feature Tree，中间是 3D 视图，右侧是属性面板，下方显示 Project 编号、在线用户和连接状态。",
    )
    add_bullets(
        doc,
        [
            "登录页：通过用户名和 roomId 进入或创建 Project。",
            "工具栏：提供 Box、Cylinder、Cone、Sphere、Sketch、Extrude、Cut、Save、Load、Undo、Redo、Move、Rotate、Scale 和 Smart Assist 入口。",
            "Feature Tree：显示建模历史与对象列表，支持选择当前编辑对象。",
            "属性面板：支持修改尺寸、半径、高度、位置、旋转和缩放参数。",
            "状态区：显示当前 Project、在线用户列表和连接状态。",
        ],
    )
    add_figure(doc, FIGURES / "cad-workspace-demo.png", "图 2-2 前端 CAD 主工作界面", width_inches=5.4)
    add_heading_2(doc, "2.3 参数化实体建模功能")
    add_body(
        doc,
        "当前课程版前端支持四类主要参数化实体：立方体（长、宽、高参数）、圆柱体（半径和高度参数）、圆锥体（底面半径、底面分段数和高度参数）以及球体（半径参数）。每个实体在创建时都会生成唯一 ID、类型、参数、创建者和创建时间，并自动加入当前 Project 的对象集合与特征树历史中。",
    )
    add_body(
        doc,
        "其中，本系统核心演示对象可聚焦为立方体、圆锥体和球体，圆柱体作为额外保留的基础实体类型一起纳入统一参数化建模流程。这样的设计既覆盖了老师关注的参数化实体类型，也保持了前端实现的统一性和可维护性。",
    )
    add_heading_2(doc, "2.4 浏览器渲染显示与视图交互")
    add_body(
        doc,
        "三维视图基于 Three.js 实现，负责创建场景、相机、灯光、辅助网格与对象网格。系统通过 OrbitControls 支持模型视图旋转、缩放与视角浏览，同时通过 TransformControls 为当前选中实体提供 Move、Rotate、Scale 三种变换模式，从而实现更直观的浏览器内交互展示。",
    )
    add_body(
        doc,
        "因此，当前前端的“旋转、拉伸、移动”能力主要体现为：一是视图层的旋转和缩放；二是 TransformControls 提供的实体拖动、旋转和缩放；三是草图到实体的拉伸生成；四是属性面板中的精确参数编辑。这样既与当前实现一致，也足以支持课程答辩演示。",
    )
    add_heading_2(doc, "2.5 草图、拉伸与对象编辑")
    add_body(
        doc,
        "系统支持 Line、Circle、Rectangle 三类草图对象，其中 Rectangle 和 Circle 可以作为 Extrude 的输入。前端会根据草图参数生成 THREE.Shape，再通过 THREE.ExtrudeGeometry 构造三维几何体，并把结果加入当前 Project。为了保证演示稳定性，Cut 操作采用了可视化标记的简化实现，不强行实现工业级布尔差集。",
    )
    add_figure(doc, FIGURES / "extrude-proof.png", "图 2-3 草图拉伸生成实体效果", width_inches=5.2)
    add_heading_2(doc, "2.6 前端协同状态管理")
    add_body(
        doc,
        "前端使用 Zustand 统一维护 Project 状态、对象集合、历史栈、在线用户列表、远端光标和连接状态。协同客户端负责与服务端建立 WebSocket 连接，发送 join-room、cursor 和 operation 消息，并接收远端 add、update、delete、replace-state 等广播。",
    )
    add_body(
        doc,
        "考虑课程设计的交付复杂度，系统采用最后写入优先（Last Write Wins）的冲突处理策略。该策略虽然没有 CRDT 的细粒度冲突合并能力，但实现简单、演示直观，足以支撑当前双窗口同 Project 协同场景。",
    )

    add_heading_1(doc, "3 后端系统设计")
    add_heading_2(doc, "3.1 后端总体结构")
    add_body(
        doc,
        "后端采用 Node.js + Express + ws 实现，不强行拆成微服务。对本课程设计而言，单体协同服务已经能够完整承担 Project 状态维护、数据保存和多客户端同步职责，也更便于在本地开发环境中快速启动与调试。",
    )
    add_bullets(
        doc,
        [
            "HTTP API：提供 Project 加载、保存、版本列表和版本恢复接口。",
            "WebSocket Hub：处理客户端连接、Project 加入和消息分发。",
            "Room State Manager：维护当前 Project 的实体对象、在线用户与光标状态。",
            "Persistence 模块：负责读写 JSON 主文件与版本快照。",
        ],
    )
    add_heading_2(doc, "3.2 Project 与实体数据文件保存")
    add_body(
        doc,
        "后端把 roomId 对应为一个 Project，并将其当前状态保存到 backend/data/projects/{roomId}.json。每个实体对象都带有唯一 ID、类型、参数、创建者、创建时间和更新时间，因此后端保存的不只是一个界面快照，而是一份可恢复的 Project 数据文件。",
    )
    add_body(
        doc,
        "用户每次点击 Save 时，后端都会额外在 backend/data/projects/{roomId}/versions/{timestamp}.json 下生成一个版本快照。这样既满足课程设计里“云上保存项目和新建实体数据文件”的要求，也为历史回溯和版本恢复提供了依据。",
    )
    add_figure(doc, FIGURES / "version-proof.png", "图 3-1 Save 与版本快照验证结果", width_inches=5.2)
    add_heading_2(doc, "3.3 多客户端协同设计")
    add_body(
        doc,
        "当多个客户端进入同一 Project 后，后端会把它们加入同一个 WebSocket 房间，并维护在线用户列表。某个客户端新增实体、修改参数或删除对象后，会以 operation 消息的形式广播给其他客户端，从而实现多客户端下的协同设计。",
    )
    add_body(
        doc,
        "与此同时，后端还会同步 presence 和 cursor 状态，用于更新在线用户列表和远端光标显示。当前版本的冲突处理采用最后写入优先策略，适合课程版的演示环境。",
    )
    add_figure(doc, FIGURES / "collab-sync-proof.png", "图 3-2 多客户端协同同步效果", width_inches=5.2)
    add_heading_2(doc, "3.4 保存、加载与版本恢复")
    add_body(
        doc,
        "后端除了保存当前 Project 状态外，还提供版本列表查询和版本恢复接口。前端可以主动请求最近一次保存状态，也可以获取历史版本列表并恢复某个版本。这使系统在课程答辩中不只是“能协同”，而是具备基本的工程管理闭环。",
    )
    add_heading_2(doc, "3.5 后端实现边界与扩展说明")
    add_body(
        doc,
        "老师模板中提到的微服务、数据库、缓存和格式转换 Worker 等内容，在本报告中不被当作硬性实现。当前课程版后端已经通过单体 Node.js 服务完成了项目保存和多客户端协同的核心目标。后续如需扩展，可再引入 PostgreSQL、Redis、对象存储或独立 Worker，但这些都属于系统演进方向，而不是本次课程设计的必做内容。",
    )

    add_heading_1(doc, "4 系统测试与成果演示")
    add_heading_2(doc, "4.1 本地运行与功能验证")
    add_body(
        doc,
        "系统要求 Node.js 18 及以上环境。前端执行 npm install 和 npm run dev 后，可在 http://localhost:5173 启动；后端执行 npm install 和 npm run dev 后，可在 http://localhost:3001 启动，WebSocket 地址为 ws://localhost:3001/ws。整个系统不依赖 Docker，适合在 macOS 本地环境中直接运行和演示。",
    )
    add_heading_2(doc, "4.2 关键功能测试项")
    add_table(
        doc,
        [
            ["测试项", "验证方式", "结果"],
            ["Project 创建", "输入 roomId 后进入新 Project", "可创建或进入协同项目"],
            ["参数化实体建模", "创建 Box、Cylinder、Cone、Sphere", "对象可正常生成"],
            ["浏览器渲染与交互", "旋转视角、拖动实体、缩放实体并修改参数", "显示与编辑正常"],
            ["草图拉伸", "对 Rectangle 或 Circle 执行 Extrude", "可生成 3D 实体"],
            ["多客户端协同", "双窗口进入同一 Project", "新增与修改实时同步"],
            ["项目保存与版本快照", "执行 Save 后检查后端文件", "生成 JSON 主文件和版本文件"],
            ["Undo / Redo", "本地历史回退与重做", "基本可用"],
        ],
        [3.6, 6.2, 6.9],
    )
    add_heading_2(doc, "4.3 成果演示截图")
    add_body(
        doc,
        "从演示结果来看，系统已经能够稳定展示前端建模、浏览器渲染、多人协同和后端版本管理等核心能力。下面给出课程答辩中最关键的几组截图。",
    )
    add_figure(doc, FIGURES / "cad-workspace-demo.png", "图 4-1 前端 CAD 主界面", width_inches=5.3)
    add_figure(doc, FIGURES / "collab-sync-proof.png", "图 4-2 双窗口同 Project 协同同步效果", width_inches=5.0)
    add_figure(doc, FIGURES / "extrude-proof.png", "图 4-3 草图拉伸效果", width_inches=5.0)
    add_figure(doc, FIGURES / "smart-assist-proof.png", "图 4-4 智能辅助建模入口", width_inches=5.0)
    add_heading_2(doc, "4.4 结果分析")
    add_body(
        doc,
        "综合验证结果可以看出，系统已经具备课程设计所要求的核心演示能力：前端能完成 Project 创建、参数化实体建模、浏览器渲染显示、实体拖动旋转缩放与草图拉伸；后端能完成项目文件保存、实体数据管理和多客户端协同同步。当前仍属于课程版 MVP 的部分主要有：Cut 为简化实现、未接入工业级几何内核、未做复杂格式导入导出、后端保持单体服务结构。",
    )

    add_heading_1(doc, "5 总结与展望", page_break=False)
    add_heading_2(doc, "5.1 总结")
    add_body(
        doc,
        "本文按照“总体架构设计、前端设计、后端设计”三条主线，完成了一个基于云协同思想的机械 CAD MVP。系统在浏览器中实现了 Project 创建、基础参数化实体建模、三维渲染显示、实体拖动旋转缩放、草图拉伸和对象参数编辑，在后端实现了项目数据文件保存、版本快照和多客户端协同设计，达到了课程设计对稳定演示和清晰讲解的要求。",
    )
    add_heading_2(doc, "5.2 后续扩展")
    add_body(
        doc,
        "后续可以在三个方向继续扩展：第一，引入 OpenCASCADE/Wasm 增强工业级几何建模能力；第二，引入 PostgreSQL、Redis 和对象存储增强后端数据底座；第三，继续扩展参数化实体类型和更复杂的前端交互能力。在当前阶段，这些内容均作为扩展方向保留，不影响本课程设计成品的完整性。",
    )

    add_heading_1(doc, "6 项目交付成果清单", page_break=False)
    add_bullets(
        doc,
        [
            "前端完整源代码：React + TypeScript + Vite + Three.js 项目。",
            "后端完整源代码：Node.js + Express + ws 协同服务与本地 JSON 持久化实现。",
            "项目文档：README、architecture、api、local-run、report-outline 等说明文档。",
            "课程报告材料：LaTeX 版报告、Word 模板完成版报告、测试矩阵与最终验收说明。",
            "系统截图与图示：架构图、UML 图、协同截图、版本管理截图、智能辅助截图。",
            "演示数据与版本快照：backend/data/projects 下的房间状态与历史版本文件。",
        ],
    )

    add_heading_1(doc, "参考文献")
    add_reference_list(
        doc,
        [
            "React Documentation. https://react.dev/.",
            "Three.js Documentation. https://threejs.org/docs/.",
            "Node.js Documentation. https://nodejs.org/docs/latest/api/.",
            "Fette I, Melnikov A. The WebSocket Protocol: RFC 6455[S]. IETF, 2011.",
            "Open CASCADE Technology Documentation. https://dev.opencascade.org/doc/overview/html/.",
        ],
    )

    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_report()
    print(output)
